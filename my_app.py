from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


Document = Document()

# profile picture
Document.add_picture('communism.jpeg', width=Inches(1.4))

# name phone number and email details
name = input('what is your name? ')
speak('hello' + name + 'how are you today? ')

phone_number = input('what is your phone number? ')
email = input('what is your email? ')

Document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
Document.add_heading('about me')
about_me = input('tell me about yourself ')
Document.add_paragraph(about_me)

# work experience
Document.add_heading('work experience')
p = Document.add_paragraph()

company = input('enter company ')
from_date = input('from date ')
to_date = input('to date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '_' + to_date + '\n').italic

experience_details = input('describe your experience at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'do you have more experiences? yes or no ')
    if has_more_experiences.lower() == 'yes':
        p = Document.add_paragraph()

        company = input('enter company ')
        from_date = input('from date ')
        to_date = input('to date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '_' + to_date + '\n').italic

        experience_details = input('describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

# skills
Document.add_heading('Skills')
skill = input('Enter skill')
p = Document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('do you have more skills? yes or no')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill')
        p = Document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break




Document.save('cv.docx')